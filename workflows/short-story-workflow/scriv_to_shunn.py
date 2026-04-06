#!/usr/bin/env python3
"""
scriv_to_shunn.py
─────────────────
Takes a compiled Scrivener manuscript (text, .docx, .rtf, or reads directly
from a .scriv project) and formats it in standard Shunn short story
manuscript format, outputting both .docx and .pdf.

Output goes to: ~/Documents/Writing/Stories/<Title>-<mm-dd-yyyy>/

Shunn short story format:
  - Times New Roman 12pt, double-spaced
  - 1-inch margins on all sides
  - First page: contact info upper-left, word count upper-right,
    title centered ~1/3 down, "by" + author name below, then story begins
  - Running header on pages 2+: "Last Name / TITLE KEYWORD / page#" right-aligned
  - 0.5-inch first-line paragraph indent
  - Scene breaks: "#" centered on its own line
  - "# # # END # # #" centered after the last paragraph

Usage:
    python scriv_to_shunn.py compiled.txt
    python scriv_to_shunn.py compiled.txt --author "Billy Gong" --title "The River"
    python scriv_to_shunn.py "My Project.scriv"
    python scriv_to_shunn.py compiled.docx --email "writer@example.com"

Requirements:
    pip install reportlab pypdf python-docx
    npm install -g docx  (or local install)
"""

import argparse
import json
import math
import os
import re
import subprocess
import sys
import tempfile
from datetime import datetime
from pathlib import Path


# ═══════════════════════════════════════════════════════════════════════
#  TEXT EXTRACTION — read compiled manuscript from various formats
# ═══════════════════════════════════════════════════════════════════════

def read_compiled_text(filepath):
    """Read compiled manuscript text from .txt, .docx, .rtf, or .scriv."""
    ext = Path(filepath).suffix.lower()

    if ext == ".txt" or ext == ".md":
        with open(filepath, "r", encoding="utf-8", errors="replace") as f:
            return f.read()

    elif ext == ".docx":
        try:
            from docx import Document
        except ImportError:
            print("ERROR: python-docx not installed. Run: pip install python-docx")
            sys.exit(1)
        doc = Document(filepath)
        return "\n".join(para.text for para in doc.paragraphs)

    elif ext == ".rtf":
        try:
            from striprtf.striprtf import rtf_to_text
        except ImportError:
            print("ERROR: striprtf not installed. Run: pip install striprtf")
            sys.exit(1)
        with open(filepath, "r", encoding="utf-8", errors="replace") as f:
            return rtf_to_text(f.read())

    elif ext == ".scriv":
        return _read_scriv_manuscript(filepath)

    else:
        print(f"ERROR: Unsupported file type '{ext}'. Use .txt, .docx, .rtf, or .scriv")
        sys.exit(1)


def _read_scriv_manuscript(scriv_path):
    """Read all manuscript text documents from a .scriv project, in binder order."""
    import xml.etree.ElementTree as ET

    scrivx_files = list(Path(scriv_path).glob("*.scrivx"))
    if not scrivx_files:
        print(f"ERROR: No .scrivx file found in {scriv_path}")
        sys.exit(1)

    tree = ET.parse(scrivx_files[0])
    root = tree.getroot()
    binder = root.find("Binder")

    # Find the DraftFolder (Manuscript)
    draft = None
    for item in binder.findall("BinderItem"):
        if item.get("Type") == "DraftFolder":
            draft = item
            break

    if draft is None:
        print("ERROR: No Manuscript (DraftFolder) found in .scrivx")
        sys.exit(1)

    # Collect all text UUIDs in binder order
    uuids = []
    _collect_text_uuids(draft, uuids)

    # Read RTF content from each
    texts = []
    for uid in uuids:
        rtf_path = os.path.join(scriv_path, "Files", "Data", uid, "content.rtf")
        if os.path.exists(rtf_path):
            text = _rtf_to_plain(rtf_path)
            if text.strip():
                texts.append(text.strip())

    return "\n\n".join(texts)


def _collect_text_uuids(item, uuids):
    """Recursively collect UUIDs of Text items in binder order."""
    item_type = item.get("Type", "Text")
    if item_type == "Text":
        uuids.append(item.get("UUID", ""))

    children = item.find("Children")
    if children is not None:
        for child in children.findall("BinderItem"):
            _collect_text_uuids(child, uuids)


def _rtf_to_plain(rtf_path):
    """Minimal RTF to plain text extraction."""
    with open(rtf_path, "r", encoding="utf-8", errors="replace") as f:
        content = f.read()

    body_match = re.search(r'\\f0\\fs\d+\s*\\cf0\s*(.*)', content, re.DOTALL)
    if body_match:
        text = body_match.group(1)
    else:
        text = content

    text = re.sub(r'\}+\s*$', '', text)
    text = re.sub(r'\\[a-zA-Z]+\d*\s?', '', text)
    text = text.replace('\\{', '{').replace('\\}', '}').replace('\\\\', '\\')
    text = text.replace('\\\n', '\n')
    text = re.sub(r'\n{3,}', '\n\n', text).strip()

    return text


# ═══════════════════════════════════════════════════════════════════════
#  MANUSCRIPT PARSING
# ═══════════════════════════════════════════════════════════════════════

def parse_manuscript(raw_text, title=None):
    """
    Parse compiled manuscript text into structured parts.

    Returns dict with:
        paragraphs: list of strings (body paragraphs)
        word_count: int
        title: str
    """
    # Clean up
    text = raw_text.strip()
    text = re.sub(r'\r\n', '\n', text)

    # Split into paragraphs (double newlines or single with blank lines)
    paragraphs = re.split(r'\n\s*\n', text)
    paragraphs = [p.strip() for p in paragraphs if p.strip()]

    # Word count
    word_count = len(text.split())

    return {
        "paragraphs": paragraphs,
        "word_count": word_count,
        "title": title,
    }


def _round_word_count(count):
    """Round word count to nearest 100 (under 1000) or nearest 500 (over 1000)."""
    if count < 1000:
        return int(round(count / 100.0)) * 100
    else:
        return int(round(count / 500.0)) * 500


def _title_keyword(title):
    """Extract a short keyword from the title for the running header."""
    # Use the first significant word (skip articles)
    words = title.upper().split()
    skip = {"THE", "A", "AN", "OF", "IN", "ON", "AT", "TO", "FOR", "AND", "OR"}
    for w in words:
        clean = re.sub(r'[^A-Z]', '', w)
        if clean and clean not in skip:
            return clean
    return words[0].upper() if words else "STORY"


# ═══════════════════════════════════════════════════════════════════════
#  DOCX GENERATION (via docx-js / Node.js)
# ═══════════════════════════════════════════════════════════════════════

def generate_docx(manuscript, output_path, author, email, address, phone):
    """Generate a Shunn-formatted .docx file using docx-js."""

    title = manuscript["title"]
    paragraphs = manuscript["paragraphs"]
    word_count = _round_word_count(manuscript["word_count"])
    keyword = _title_keyword(title)
    last_name = author.split()[-1] if author else "Author"

    # Build contact block lines
    contact_lines = [author]
    if address:
        contact_lines.extend(address.split("\\n"))
    if phone:
        contact_lines.append(phone)
    if email:
        contact_lines.append(email)

    # Escape strings for JS
    def js_str(s):
        return json.dumps(s)

    # Build paragraph JS entries
    body_paragraphs_js = []
    for para in paragraphs:
        # Detect scene breaks
        stripped = para.strip()
        if stripped in ("#", "* * *", "***", "---", "– – –", "• • •"):
            body_paragraphs_js.append(f"""
    new Paragraph({{
      alignment: AlignmentType.CENTER,
      spacing: {{ before: 240, after: 240, line: 480 }},
      children: [new TextRun({{ text: "#", font: "Times New Roman", size: 24 }})]
    }})""")
        else:
            # Escape for JS string
            escaped = para.replace("\\", "\\\\").replace('"', '\\"').replace("\n", " ")
            body_paragraphs_js.append(f"""
    new Paragraph({{
      spacing: {{ line: 480 }},
      indent: {{ firstLine: 720 }},
      children: [new TextRun({{ text: "{escaped}", font: "Times New Roman", size: 24 }})]
    }})""")

    body_js = ",".join(body_paragraphs_js)

    # Build contact block paragraphs
    contact_js_parts = []
    for line in contact_lines:
        escaped = line.replace("\\", "\\\\").replace('"', '\\"')
        contact_js_parts.append(f"""
    new Paragraph({{
      spacing: {{ line: 240 }},
      children: [new TextRun({{ text: "{escaped}", font: "Times New Roman", size: 24 }})]
    }})""")
    contact_js = ",".join(contact_js_parts)

    js_code = f"""
const fs = require("fs");
const {{ Document, Packer, Paragraph, TextRun, Header, Footer,
         AlignmentType, PageNumber, PageBreak, TabStopType, TabStopPosition }} = require("docx");

const doc = new Document({{
  styles: {{
    default: {{
      document: {{
        run: {{ font: "Times New Roman", size: 24 }}
      }}
    }}
  }},
  sections: [
    // ── FIRST PAGE (no header) ──
    {{
      properties: {{
        page: {{
          size: {{ width: 12240, height: 15840 }},
          margin: {{ top: 1440, right: 1440, bottom: 1440, left: 1440 }}
        }},
        titlePage: true
      }},
      headers: {{
        default: new Header({{
          children: [new Paragraph({{
            alignment: AlignmentType.RIGHT,
            children: [
              new TextRun({{ text: "{last_name} / {keyword} / ", font: "Times New Roman", size: 24 }}),
              new TextRun({{ children: [PageNumber.CURRENT], font: "Times New Roman", size: 24 }})
            ]
          }})]
        }}),
        first: new Header({{ children: [new Paragraph("")] }})
      }},
      children: [
        // Contact info (upper left)
        {contact_js},

        // Word count (right-aligned, after contact block)
        new Paragraph({{
          alignment: AlignmentType.RIGHT,
          spacing: {{ before: 0, line: 240 }},
          children: [new TextRun({{ text: "About {word_count} words", font: "Times New Roman", size: 24 }})]
        }}),

        // Vertical space before title (~1/3 down page)
        new Paragraph({{ spacing: {{ before: 2400 }}, children: [] }}),

        // Title
        new Paragraph({{
          alignment: AlignmentType.CENTER,
          spacing: {{ after: 120, line: 240 }},
          children: [new TextRun({{ text: {js_str(title)}, font: "Times New Roman", size: 24, bold: false }})]
        }}),

        // "by"
        new Paragraph({{
          alignment: AlignmentType.CENTER,
          spacing: {{ after: 120, line: 240 }},
          children: [new TextRun({{ text: "by", font: "Times New Roman", size: 24 }})]
        }}),

        // Author name
        new Paragraph({{
          alignment: AlignmentType.CENTER,
          spacing: {{ after: 480, line: 240 }},
          children: [new TextRun({{ text: {js_str(author)}, font: "Times New Roman", size: 24 }})]
        }}),

        // Body paragraphs
        {body_js},

        // END marker
        new Paragraph({{
          alignment: AlignmentType.CENTER,
          spacing: {{ before: 480, line: 480 }},
          children: [new TextRun({{ text: "# # # END # # #", font: "Times New Roman", size: 24 }})]
        }})
      ]
    }}
  ]
}});

Packer.toBuffer(doc).then(buffer => {{
  fs.writeFileSync({js_str(str(output_path))}, buffer);
  console.log("DOCX created: " + {js_str(str(output_path))});
}});
"""

    # Write and run the JS
    with tempfile.NamedTemporaryFile(mode="w", suffix=".js", delete=False, dir="/tmp") as f:
        f.write(js_code)
        js_path = f.name

    try:
        env = os.environ.copy()
        npm_global = "/sessions/charming-ecstatic-albattani/.npm-global/lib/node_modules"
        if os.path.exists(npm_global):
            env["NODE_PATH"] = npm_global
        result = subprocess.run(
            ["node", js_path],
            capture_output=True, text=True, env=env, timeout=30
        )
        if result.returncode != 0:
            print(f"ERROR generating .docx:\n{result.stderr}")
            sys.exit(1)
        print(result.stdout.strip())
    finally:
        os.unlink(js_path)


# ═══════════════════════════════════════════════════════════════════════
#  PDF GENERATION (via reportlab)
# ═══════════════════════════════════════════════════════════════════════

def generate_pdf(manuscript, output_path, author, email, address, phone):
    """Generate a Shunn-formatted .pdf using reportlab."""
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.units import inch
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_RIGHT
    from reportlab.platypus import (
        BaseDocTemplate, PageTemplate, Frame, Paragraph, Spacer,
        NextPageTemplate, PageBreak
    )

    title = manuscript["title"]
    paragraphs = manuscript["paragraphs"]
    word_count = _round_word_count(manuscript["word_count"])
    keyword = _title_keyword(title)
    last_name = author.split()[-1] if author else "Author"

    PAGE_W, PAGE_H = letter  # 8.5 x 11 inches
    MARGIN = 1 * inch

    # ── Styles ──
    style_normal = ParagraphStyle(
        "ShuBody",
        fontName="Times-Roman",
        fontSize=12,
        leading=24,  # double-spaced (12pt * 2)
        firstLineIndent=0.5 * inch,
        spaceBefore=0,
        spaceAfter=0,
    )
    style_contact = ParagraphStyle(
        "ShuContact",
        fontName="Times-Roman",
        fontSize=12,
        leading=14,
        alignment=TA_LEFT,
    )
    style_wordcount = ParagraphStyle(
        "ShuWordCount",
        fontName="Times-Roman",
        fontSize=12,
        leading=14,
        alignment=TA_RIGHT,
    )
    style_title = ParagraphStyle(
        "ShuTitle",
        fontName="Times-Roman",
        fontSize=12,
        leading=14,
        alignment=TA_CENTER,
    )
    style_scene_break = ParagraphStyle(
        "ShuSceneBreak",
        fontName="Times-Roman",
        fontSize=12,
        leading=24,
        alignment=TA_CENTER,
        spaceBefore=12,
        spaceAfter=12,
    )
    style_end = ParagraphStyle(
        "ShuEnd",
        fontName="Times-Roman",
        fontSize=12,
        leading=24,
        alignment=TA_CENTER,
        spaceBefore=24,
    )

    # ── Header callback ──
    def _header_later_pages(canvas_obj, doc):
        canvas_obj.saveState()
        canvas_obj.setFont("Times-Roman", 12)
        header_text = f"{last_name} / {keyword} / {doc.page}"
        canvas_obj.drawRightString(PAGE_W - MARGIN, PAGE_H - 0.5 * inch, header_text)
        canvas_obj.restoreState()

    def _header_first_page(canvas_obj, doc):
        pass  # No header on first page

    # ── Build document ──
    doc = BaseDocTemplate(
        str(output_path),
        pagesize=letter,
        leftMargin=MARGIN,
        rightMargin=MARGIN,
        topMargin=MARGIN,
        bottomMargin=MARGIN,
    )

    frame = Frame(MARGIN, MARGIN, PAGE_W - 2 * MARGIN, PAGE_H - 2 * MARGIN, id="main")

    doc.addPageTemplates([
        PageTemplate(id="first", frames=[frame], onPage=_header_first_page),
        PageTemplate(id="later", frames=[frame], onPage=_header_later_pages),
    ])

    # ── Story elements ──
    story = []

    # Contact info
    contact_lines = [author]
    if address:
        contact_lines.extend(address.split("\n"))
    if phone:
        contact_lines.append(phone)
    if email:
        contact_lines.append(email)

    for line in contact_lines:
        story.append(Paragraph(line, style_contact))

    # Word count (right-aligned)
    story.append(Spacer(1, 6))
    story.append(Paragraph(f"About {word_count} words", style_wordcount))

    # Vertical space (~1/3 down page)
    story.append(Spacer(1, 2 * inch))

    # Title block
    story.append(Paragraph(title, style_title))
    story.append(Spacer(1, 6))
    story.append(Paragraph("by", style_title))
    story.append(Spacer(1, 6))
    story.append(Paragraph(author, style_title))
    story.append(Spacer(1, 0.5 * inch))

    # Switch to "later" template after first page
    story.append(NextPageTemplate("later"))

    # Body paragraphs
    for para in paragraphs:
        stripped = para.strip()
        # Detect scene breaks
        if stripped in ("#", "* * *", "***", "---", "– – –", "• • •"):
            story.append(Paragraph("#", style_scene_break))
        else:
            # Escape HTML entities for reportlab
            safe = (para
                    .replace("&", "&amp;")
                    .replace("<", "&lt;")
                    .replace(">", "&gt;")
                    .replace("\n", " "))
            story.append(Paragraph(safe, style_normal))

    # END marker
    story.append(Paragraph("# # # END # # #", style_end))

    doc.build(story)
    print(f"PDF created: {output_path}")


# ═══════════════════════════════════════════════════════════════════════
#  COVER LETTER GENERATION (via docx-js / Node.js)
# ═══════════════════════════════════════════════════════════════════════

def generate_cover_letter(manuscript, output_path, author, email, address, phone):
    """Generate a standard fiction submission cover letter as .docx."""

    title = manuscript["title"]
    word_count = _round_word_count(manuscript["word_count"])
    today = datetime.now().strftime("%B %d, %Y")  # e.g. "April 03, 2026"

    # Escape strings for JS
    def js_str(s):
        return json.dumps(s)

    # Build contact block for closing
    contact_lines = []
    if address:
        contact_lines.extend(address.split("\\n"))
    if phone:
        contact_lines.append(phone)
    if email:
        contact_lines.append(email)

    contact_closing_js = ""
    for line in contact_lines:
        escaped = line.replace("\\", "\\\\").replace('"', '\\"')
        contact_closing_js += f"""
    new Paragraph({{
      spacing: {{ line: 276 }},
      children: [new TextRun({{ text: "{escaped}", font: "Times New Roman", size: 24 }})]
    }}),"""

    js_code = f"""
const fs = require("fs");
const {{ Document, Packer, Paragraph, TextRun, AlignmentType, TabStopType, TabStopPosition }} = require("docx");

const doc = new Document({{
  styles: {{
    default: {{
      document: {{
        run: {{ font: "Times New Roman", size: 24 }}
      }}
    }}
  }},
  sections: [{{
    properties: {{
      page: {{
        size: {{ width: 12240, height: 15840 }},
        margin: {{ top: 1440, right: 1440, bottom: 1440, left: 1440 }}
      }}
    }},
    children: [
      // Date
      new Paragraph({{
        spacing: {{ after: 240, line: 276 }},
        children: [new TextRun({{ text: {js_str(today)}, font: "Times New Roman", size: 24 }})]
      }}),

      // Editor placeholder
      new Paragraph({{
        spacing: {{ line: 276 }},
        children: [new TextRun({{ text: "[Editor Name]", font: "Times New Roman", size: 24, italics: true, color: "888888" }})]
      }}),
      new Paragraph({{
        spacing: {{ line: 276 }},
        children: [new TextRun({{ text: "[Publication Name]", font: "Times New Roman", size: 24, italics: true, color: "888888" }})]
      }}),
      new Paragraph({{
        spacing: {{ after: 240, line: 276 }},
        children: [new TextRun({{ text: "[Address]", font: "Times New Roman", size: 24, italics: true, color: "888888" }})]
      }}),

      // Salutation
      new Paragraph({{
        spacing: {{ after: 240, line: 276 }},
        children: [new TextRun({{ text: "Dear [Editor Name],", font: "Times New Roman", size: 24 }})]
      }}),

      // Body paragraph 1 — the pitch
      new Paragraph({{
        spacing: {{ after: 240, line: 276 }},
        children: [new TextRun({{
          text: "Please find enclosed my short story, ",
          font: "Times New Roman", size: 24
        }}), new TextRun({{
          text: {js_str('"' + title + '"')},
          font: "Times New Roman", size: 24
        }}), new TextRun({{
          text: ", at approximately {word_count} words. [One or two sentences describing the story — its premise, central tension, or what makes it compelling. Keep it brief and let the story speak for itself.]",
          font: "Times New Roman", size: 24
        }})]
      }}),

      // Body paragraph 2 — bio placeholder
      new Paragraph({{
        spacing: {{ after: 240, line: 276 }},
        children: [new TextRun({{
          text: "[Author bio — mention relevant publications, awards, or professional background. If this is your first submission, you can omit this paragraph or simply state that this is your first publication.]",
          font: "Times New Roman", size: 24, italics: true, color: "888888"
        }})]
      }}),

      // Body paragraph 3 — simultaneous submission / closing
      new Paragraph({{
        spacing: {{ after: 240, line: 276 }},
        children: [new TextRun({{
          text: "This is a simultaneous submission. Thank you for your time and consideration.",
          font: "Times New Roman", size: 24
        }})]
      }}),

      // Closing
      new Paragraph({{
        spacing: {{ after: 60, line: 276 }},
        children: [new TextRun({{ text: "Sincerely,", font: "Times New Roman", size: 24 }})]
      }}),

      // Author name
      new Paragraph({{
        spacing: {{ after: 120, line: 276 }},
        children: [new TextRun({{ text: {js_str(author)}, font: "Times New Roman", size: 24 }})]
      }}),

      // Contact info
      {contact_closing_js}
    ]
  }}]
}});

Packer.toBuffer(doc).then(buffer => {{
  fs.writeFileSync({js_str(str(output_path))}, buffer);
  console.log("Cover letter created: " + {js_str(str(output_path))});
}});
"""

    # Write and run the JS
    with tempfile.NamedTemporaryFile(mode="w", suffix=".js", delete=False, dir="/tmp") as f:
        f.write(js_code)
        js_path = f.name

    try:
        env = os.environ.copy()
        npm_global = "/sessions/charming-ecstatic-albattani/.npm-global/lib/node_modules"
        if os.path.exists(npm_global):
            env["NODE_PATH"] = npm_global
        result = subprocess.run(
            ["node", js_path],
            capture_output=True, text=True, env=env, timeout=30
        )
        if result.returncode != 0:
            print(f"ERROR generating cover letter:\n{result.stderr}")
            sys.exit(1)
        print(result.stdout.strip())
    finally:
        os.unlink(js_path)


# ═══════════════════════════════════════════════════════════════════════
#  MAIN WORKFLOW
# ═══════════════════════════════════════════════════════════════════════

def run(input_path, title=None, author="W. S. Gong", email="",
        address="", phone="", output_dir_override=None):
    """
    Full pipeline: read compiled text → format as Shunn → output .docx + .pdf.

    Returns dict with output_dir, docx_path, pdf_path.
    """
    # Read source
    print(f"Reading: {input_path}")
    raw_text = read_compiled_text(input_path)

    if not raw_text.strip():
        print("ERROR: No text found in input.")
        sys.exit(1)

    # Determine title
    if not title:
        title = Path(input_path).stem
        # Clean up common suffixes
        for suffix in ("-compiled", "_compiled", "-export", "_export", " Compiled"):
            title = title.replace(suffix, "")
        title = title.strip()

    # Parse manuscript
    manuscript = parse_manuscript(raw_text, title)
    print(f"  Title: {title}")
    print(f"  Words: {manuscript['word_count']} (≈{_round_word_count(manuscript['word_count'])})")
    print(f"  Paragraphs: {len(manuscript['paragraphs'])}")

    # Create output directory
    date_str = datetime.now().strftime("%m-%d-%Y")
    safe_title = re.sub(r'[^\w\s-]', '', title).strip()
    safe_title = re.sub(r'\s+', ' ', safe_title)
    dir_name = f"{safe_title}-{date_str}"

    if output_dir_override:
        output_dir = os.path.join(output_dir_override, dir_name)
    else:
        output_dir = os.path.expanduser(f"~/Documents/Writing/Stories/{dir_name}")

    os.makedirs(output_dir, exist_ok=True)
    print(f"  Output: {output_dir}")

    # Generate .docx
    docx_path = os.path.join(output_dir, f"{safe_title}.docx")
    print("\nGenerating .docx...")
    generate_docx(manuscript, docx_path, author, email, address, phone)

    # Generate .pdf
    pdf_path = os.path.join(output_dir, f"{safe_title}.pdf")
    print("Generating .pdf...")
    generate_pdf(manuscript, pdf_path, author, email, address, phone)

    # Generate cover letter
    cover_path = os.path.join(output_dir, f"Cover Letter - {safe_title}.docx")
    print("Generating cover letter...")
    generate_cover_letter(manuscript, cover_path, author, email, address, phone)

    print(f"\n{'─' * 50}")
    print(f"✓ Shunn manuscript formatted!")
    print(f"  Directory: {output_dir}")
    print(f"  DOCX:      {os.path.basename(docx_path)}")
    print(f"  PDF:       {os.path.basename(pdf_path)}")
    print(f"  Cover:     {os.path.basename(cover_path)}")
    print(f"{'─' * 50}")

    return {
        "output_dir": output_dir,
        "docx_path": docx_path,
        "pdf_path": pdf_path,
        "cover_path": cover_path,
    }


# ═══════════════════════════════════════════════════════════════════════
#  CLI
# ═══════════════════════════════════════════════════════════════════════

def main():
    parser = argparse.ArgumentParser(
        description="Format a compiled Scrivener manuscript in Shunn short story format (.docx + .pdf).",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Shunn format includes:
  - Times New Roman 12pt, double-spaced, 1-inch margins
  - Contact info, word count, title page
  - Running header: "Last Name / KEYWORD / page#"
  - Scene breaks as centered "#"
  - "# # # END # # #" marker

Examples:
  %(prog)s compiled.txt --author "Billy Gong" --title "The River"
  %(prog)s compiled.docx --email "billy@example.com"
  %(prog)s "My Project.scriv" --author "Billy Gong"
  %(prog)s story.txt --output-dir /tmp/test/
        """,
    )

    parser.add_argument(
        "input",
        help="Compiled manuscript (.txt, .docx, .rtf) or .scriv project",
    )
    parser.add_argument(
        "--title", "-t",
        default=None,
        help="Story title (default: derived from filename)",
    )
    parser.add_argument(
        "--author", "-a",
        default="W. S. Gong",
        help="Author name (default: W. S. Gong)",
    )
    parser.add_argument(
        "--email", "-e",
        default="",
        help="Contact email for the title page",
    )
    parser.add_argument(
        "--address",
        default="",
        help="Mailing address for the title page (use \\n for line breaks)",
    )
    parser.add_argument(
        "--phone",
        default="",
        help="Phone number for the title page",
    )
    parser.add_argument(
        "--output-dir",
        default=None,
        help="Override output directory (default: ~/Documents/Writing/Stories/<title>-<date>)",
    )

    args = parser.parse_args()

    result = run(
        input_path=args.input,
        title=args.title,
        author=args.author,
        email=args.email,
        address=args.address,
        phone=args.phone,
        output_dir_override=args.output_dir,
    )


if __name__ == "__main__":
    main()
